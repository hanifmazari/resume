import io
import re
import json
import time
import os
from dotenv import load_dotenv
from groq import Groq
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document

load_dotenv()

# --- API KEY POOL MANAGEMENT ---

def load_groq_api_keys():
    keys = []
    i = 1
    while True:
        key = os.getenv(f'GROQ_API_KEY{i}')
        if not key:
            break
        keys.append(key)
        i += 1
    return keys

class GroqClientPool:
    def __init__(self, api_keys):
        if not api_keys:
            raise Exception("No GROQ_API_KEY# variables found in environment.")
        self.api_keys = api_keys
        self.clients = [Groq(api_key=key) for key in api_keys]
        self.current_index = 0
    
    def get_current_client(self):
        return self.clients[self.current_index]
    
    def switch_to_next_client(self):
        self.current_index = (self.current_index + 1) % len(self.clients)
        print(f"Switching to next API key: index {self.current_index}")
    
    def chat_completion(self, **kwargs):
        attempts = 0
        max_attempts = len(self.clients)
        while attempts < max_attempts:
            client = self.get_current_client()
            try:
                completion = client.chat.completions.create(**kwargs)
                return completion
            except Exception as e:
                # You can add more specific error handling if needed (e.g., check for rate limit error code)
                print(f"API key index {self.current_index} failed with error: {e}")
                self.switch_to_next_client()
                attempts += 1
                time.sleep(1)  # brief pause before retrying
        raise Exception("All API keys exhausted or failed.")

# Initialize client pool
api_keys = load_groq_api_keys()
client_pool = GroqClientPool(api_keys)

# --- FILE TEXT EXTRACTION FUNCTIONS ---

def detect_file_type(file_bytes):
    if file_bytes.startswith(b'%PDF'):
        return 'pdf'
    if file_bytes.startswith(b'PK\x03\x04'):
        try:
            Document(io.BytesIO(file_bytes))
            return 'docx'
        except:
            pass
    try:
        file_bytes.decode('utf-8')
        return 'txt'
    except UnicodeDecodeError:
        pass
    for encoding in ['latin-1', 'cp1252']:
        try:
            file_bytes.decode(encoding)
            return 'txt'
        except UnicodeDecodeError:
            continue
    return None

def extract_text_from_file(file_bytes, filename=None):
    if filename:
        ext = filename.lower().split('.')[-1]
    else:
        ext = detect_file_type(file_bytes)
        if not ext:
            return None
    
    if ext == 'pdf':
        pdf_stream = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    
    elif ext == 'docx':
        doc_stream = io.BytesIO(file_bytes)
        document = Document(doc_stream)
        text = "\n".join([para.text for para in document.paragraphs])
        return text
    
    elif ext == 'txt':
        try:
            return file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            for encoding in ['latin-1', 'cp1252', 'ascii']:
                try:
                    return file_bytes.decode(encoding)
                except UnicodeDecodeError:
                    continue
        return None
    else:
        raise ValueError(f"Unsupported file type: {ext}")

# --- JSON EXTRACTION ---

def extract_json_from_text(text):
    json_match = re.search(r'{.+}', text, re.DOTALL)
    if json_match:
        json_str = json_match.group(0)
        try:
            return json.loads(json_str)
        except:
            pass
    try:
        cleaned_text = re.sub(r'```json\s*|\s*```', '', text).strip()
        return json.loads(cleaned_text)
    except:
        return None

# --- GROQ API CALLS USING POOL ---

def get_groq_response(prompt):
    try:
        completion = client_pool.chat_completion(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_completion_tokens=1024,
            top_p=1,
            stream=False
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        return f"Error: {str(e)}"

# --- RESUME ANALYSIS FUNCTION ---

def analyze_resume_against_jd(resume_file_bytes, job_description):
    resume_text = extract_text_from_file(resume_file_bytes)

    input_prompt = """
Act as an ATS (Application Tracking System) expert. Analyze the resume against the job description.

Resume: {text}
Job Description: {jd}

Return ONLY a JSON with:
{{"JD Match":"X%","MissingKeywords":["keyword1","keyword2"]}}
"""

    formatted_prompt = input_prompt.format(text=resume_text, jd=job_description)
    response = get_groq_response(formatted_prompt)
    result = extract_json_from_text(response)

    if result:
        match_percentage = result.get("JD Match", "0%").replace("%", "")
        try:
            match_value = float(match_percentage)
        except:
            match_value = 0

        if match_value >= 80:
            assessment = "Excellent match! Your resume is well-aligned with this job."
        elif match_value >= 60:
            assessment = "Good match. With some improvements, your resume could be great for this role."
        else:
            assessment = "Your resume needs significant improvements to match this job description."

        result["assessment"] = assessment
        return result
    else:
        return {"error": "Failed to parse results"}

# --- PROFILE, EXPERIENCE, PROJECT ENHANCERS ---

def enhance_profile_summary(summary: str) -> str:
    prompt = (
        "You are a professional resume writer. Rewrite the following resume profile summary to make it more impactful, concise, and professional. "
        "Ensure the result is a single paragraph with no bullet points, and do not include any introductory or explanatory text—return only the improved summary:\n\n"
        f"{summary}"
    )
    completion = client_pool.chat_completion(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_completion_tokens=512,
        top_p=1,
        stream=False
    )
    return completion.choices[0].message.content.strip()

def enhance_professional_experience(experience: str) -> str:
    prompt = (
        "You are a professional resume writer. Rewrite the following professional experience to make it results-oriented and impactful. "
        "Focus on achievements and measurable outcomes. Format it as a single paragraph, limited to less than 3 lines. "
        "Do not use bullet points, headers, or any extra labels—only return the enhanced paragraph:\n\n"
        f"{experience}"
    )
    completion = client_pool.chat_completion(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_completion_tokens=512,
        top_p=1,
        stream=False
    )
    return completion.choices[0].message.content.strip()

def enhance_project_description(project_desc: str) -> str:
    prompt = (
        "You are a professional resume writer. Rewrite the following project description to make it clear, results-oriented, and impactful. "
        "Highlight key achievements, technologies used, and measurable outcomes. "
        "Format it as a single paragraph, limited to less than 3 lines. "
        "Do not use bullet points, headers, or any extra labels—only return the enhanced paragraph:\n\n"
        f"{project_desc}"
    )
    completion = client_pool.chat_completion(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_completion_tokens=512,
        top_p=1,
        stream=False
    )
    return completion.choices[0].message.content.strip()